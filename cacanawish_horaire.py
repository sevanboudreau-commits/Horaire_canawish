# dependances : 
# pip install python-docx
# pip install PySide6 --no-compile --no-cache-dir

import sys
from PySide6.QtWidgets import QApplication, QMainWindow, QDialog, QLineEdit
from ui_HoraireInput import Ui_Horaire
import datetime
from docx import Document
from docx.enum.section import WD_ORIENT



class InputDialog(QDialog):

    def __init__(self):
        super().__init__()
        self.ui = Ui_Horaire()
        self.ui.setupUi(self)

        # Connect the built-in button box to close the dialog
        self.ui.buttonBox.accepted.connect(self.accept)
        self.ui.buttonBox.rejected.connect(self.reject)

    def get_N(self):
        return self.ui.spinBox.value()
    def get_monits(self):
        return self.ui.spinBox_2.value()
    def get_monits_n_dodo(self):
        return self.ui.spinBox_3.value()
    def get_n_aides(self):
        return self.ui.spinBox_4.value()
    def get_n_dortoirs(self):
        return self.ui.spinBox_5.value()
    def get_n_dodos_ETA(self):
        return self.ui.spinBox_6.value()


"""class MainWindow(QMainWindow):

    def __init__(self, num_inputs):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        # Generate the inputs immediately on creation based on the passed number
        self.generate_fields(num_inputs)

    def generate_fields(self, count):
        layout = self.ui.scrollAreaWidgetContents.layout()

        for i in range(count):
            line_edit = QLineEdit(self)
            line_edit.setPlaceholderText(f"Input field #{i+1}")
            layout.addWidget(line_edit)"""


def horaire_output(N,monits,monits_n_dodos,n_aides,n_dortoirs,n_dodos_ETA):
    """N=int(input("Nombres de dodos: "))
    monits=int(input("Nombres de moniteurs: "))
    monits_n_dodos=int(input("Nombre de moniteurs qui ne peuvent pas faire le dodo: "))
    n_aides=int(input("Nombre d'aides moniteur : "))
    f= 3 #int(input("Frequence de surveillance/sieste pour non-dodos: "))
    n_dortoirs = int(input("Nombre de dortoires: "))
    n_dodos_ETA =int(input("Nombre de dodos avant le premier dodo avec les campeurs : "))"""
    
    def iszero(i): 
        if i==0 :return 1
        else : return 0

    monits_dodos = monits-monits_n_dodos
    n_dodo_p_n=n_dortoirs*2
    n_dodos_tot = N*n_dodo_p_n
    max_de_suite= n_dodo_p_n//(monits_dodos-n_dodo_p_n)
    print(max_de_suite)

    jours ={}
    for i in range(N):
        jours[i]=[(j-1)%monits_dodos+1 for j in range(1+n_dodo_p_n*i,1+n_dodo_p_n*(i+1))]
    horaire_monits={}
    for i in range(monits_dodos):
        horaire_monits[i+1]=[(monits_dodos*j+i)//n_dodo_p_n+1 for j in range((n_dodos_tot-i-1)//monits_dodos+1)]
    
    
    """n_dodos_min = n_dodos_tot//monits_dodos
    n_dodos_sup = n_dodos_tot%monits_dodos
    n_dodos_max = n_dodos_min+1-iszero(n_dodos_sup)
    print(f"Nombre de nuits min: {n_dodos_min}")
    print(f"Nombre de nuits max: {n_dodos_max}")"""
    #print(f"La configuration avec le nombre maximal de nuits de suite a {max_de_suite} nuits de suite")
    
    # Dodos
    print(f"\nHoraires des nuits:")
    conges_am = [0]*N
    conges_pm = [0]*N
    for i in range(1,monits_dodos+1):
        temp=0
        b=True
        for j in range(N):
            if j+1 in horaire_monits[i]:
                temp+=1
                if temp == max_de_suite and conges_am[j]==0 and b:
                    b=False
                    conges_am[j]=i
                print(str(i)+" ",end="")
            else: 
                print("X ",end="")
                temp=0
        print()
    print(horaire_monits,"\n")

    # Pauses
    print("Horaires des pauses:")
    for i in range(1,monits_dodos+1): #prob inutile
        temp=0
        if i not in conges_am:
            for j in range(1,N+1):
                if j in horaire_monits[i]:
                    temp+=1
                    if conges_am[j-1]==0 and  temp==max_de_suite:
                        conges_am[j-1]=i
                        break
                else : temp =0
    for i in range(monits_dodos,0,-1):
        if i not in conges_am:
            for j in horaire_monits[i]:
                if conges_am[j-1]==0:
                    conges_am[j-1]=i
                    break
    for i in range(1,monits_dodos+1):
        if i not in conges_am:
            for j in horaire_monits[i]:
                if conges_pm[j-1]==0:
                    conges_pm[j-1]=i
                    break
    for i in range(monits_dodos+1,monits+1):
        for j in range(N):
            if conges_am[j]==0:
                conges_am[j]=i
                break
    for i in range(monits_dodos+1,monits+1):
        if i not in conges_am:
            for j in range(N):
                if conges_pm[j]==0:
                    conges_pm[j]=i
                    break
    print(conges_am)
    print(conges_pm,"\n")
    
    douches={}
    surveillances = {}
    siestes = {}
    bricos = {}
    for i in range(N):douches[i]=[]
    for i in range(N):surveillances[i]=[]
    for i in range(N-1):siestes[i]=[]
    for i in range(N):bricos[i]=[]
    monits_douches = n_dortoirs
    monits_surveillances = n_dortoirs+1
    monits_siestes = n_dortoirs+1
    monits_bricos=2
    #sieste surveillance brico douche
    print("Horaire des surveillances")
    for _ in range(1):
        for i in range(monits_dodos):
            for j in range(N-1):
                if len(surveillances[j])<monits_surveillances and (monits_dodos-i) not in jours[j] and (monits_dodos-i) not in surveillances[j]:
                    surveillances[j].append(monits_dodos-i)
                    break
    while sum([len(i) for i in surveillances.values()])< monits_surveillances*(N):
        for i in list(range(monits_dodos+1,monits+1))+list(range(monits_dodos,0,-1))+list(range(monits_dodos+1,monits+1)):
            for j in range(N):
                if len(surveillances[j])<monits_surveillances and i not in surveillances[j] and i not in jours[j]:
                    surveillances[j].append(i)
                    break
        

    print(surveillances,"\n")

    print("Horaire des siestes")
    
    for _ in range(0):
        for i in range(monits_dodos):
            for j in range(N-1):
                if len(siestes[j])<monits_siestes and (monits_dodos-i) not in siestes[j]and (monits_dodos-i) not in jours[j] :
                    siestes[j].append(monits_dodos-i)
                    break
    while sum([len(i) for i in siestes.values()])< monits_siestes*(N-1):
        for i in list(range(monits_dodos+1,monits+1))+list(range(monits_dodos,0,-1))+list(range(monits_dodos+1,monits+1)) :
            for j in range(N-1):
                if len(siestes[j])<monits_siestes and i not in siestes[j] and (i>monits-monits_n_dodos or i not in surveillances[j]) :
                    siestes[j].append(i)
                    break
    print(siestes,"\n")

    print("Horaire des Bricolages")
    while sum([len(i) for i in bricos.values()])!= monits_bricos*(N):
        for i in range(1,monits+1):
            for j in range(N):
                if len(bricos[j])<monits_bricos and conges_am[j]!= i and i not in bricos[j] :
                    bricos[j].append(i)
                    break
        
    print(bricos,"\n")

    print("Horaire des douches")
    while sum([len(i) for i in douches.values()])!= monits_douches*(N):
        for i in range(monits-n_aides):
            for j in range(N):
                if len(douches[j])<monits_douches and conges_am[j]!= (monits-i-n_aides) and (monits-i-n_aides) not in douches[j]:
                    douches[j].append(monits-i-n_aides)
                    break
    print(douches,"\n")
    print("nonbre de personnes par jour pour siestes, surveil, bricos, douches : ")
    print([len(i) for i in siestes.values()])
    print([len(i) for i in surveillances.values()])
    print([len(i) for i in bricos.values()])
    print([len(i) for i in douches.values()])
    

    # Create a brand new document object
    doc = Document()

    # 1. Grab the current section (the first one)
    section = doc.sections[0]

    # 2. Set the orientation to landscape
    section.orientation = WD_ORIENT.LANDSCAPE

    # 3. Swap the width and height measurements to match the landscape bounds
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    doc.add_heading('Horaire', level=0)
    
    table = doc.add_table(rows=1+(4*n_dortoirs)+3, cols=N+3)
    table.style = 'Light Grid Accent 1'

    now=datetime.datetime.now()
    for i in range(N+1):
        table.rows[0].cells[i+2].text= str((now+datetime.timedelta(days=i+n_dodos_ETA)).date())
    for i in range(n_dortoirs):
        if i == 0 : x="Iroquois"
        elif i == 1 : x="Jongleuse"
        elif i == 2 : x="Ruche"
        else: x="Je sais pas"
        table.rows[1+i*4].cells[0].text= x
        table.rows[1+i*4].cells[1].text= "Douche"
        for j in range(N):
            table.rows[1+i*4].cells[j+3].text=str(douches[j][i])+" "
        table.rows[2+i*4].cells[1].text= "Sieste"
        for j in range(N-1):
            if i==0 :table.rows[2+i*4].cells[j+3].text=str(siestes[j][i])+" \n"+str(siestes[j][i+1])+" "
            else :table.rows[2+i*4].cells[j+3].text=str(siestes[j][i+1])+" "
        table.rows[3+i*4].cells[1].text= "Surveillance"
        for j in range(N):
            if i==0 :table.rows[3+i*4].cells[j+2].text=str(surveillances[j][i])+" \n"+str(surveillances[j][i+1])+" "
            else :table.rows[3+i*4].cells[j+2].text=str(surveillances[j][i+1])+" "
        table.rows[4+i*4].cells[1].text= "Dodo"
        for j in range(N):
            table.rows[4+i*4].cells[j+2].text=str(jours[j][i])+" \n"+str(jours[j][i+n_dortoirs])+" "
        table.cell(i*4+1, 0).merge(table.cell(i*4+4, 0))
        
    table.rows[1+(4*n_dortoirs)].cells[0].text= "Bricos"
    for j in range(N):
        table.rows[1+(4*n_dortoirs)].cells[j+3].text=str(bricos[j][0])+" \n"+str(bricos[j][1])+" "
    table.rows[2+(4*n_dortoirs)].cells[0].text="Pause 4h"
    table.rows[2+(4*n_dortoirs)].cells[1].text="Avant-midi"
    table.rows[3+(4*n_dortoirs)].cells[1].text="Après-midi"
    for j in range(N):
        table.rows[2+(4*n_dortoirs)].cells[j+3].text=str(conges_am[j])+" "
    for j in range(N):
        table.rows[3+(4*n_dortoirs)].cells[j+3].text=str(conges_pm[j])+" "
    
    doc.save('Horaire.docx')






if __name__ == "__main__":
    app = QApplication(sys.argv)

    # 1. Create and show the Dialog first
    dialog = InputDialog()

    # exec() blocks the rest of the code until the dialog is closed.
    # It returns QDialog.DialogCode.Accepted (usually 1) if they clicked OK.
    if dialog.exec() == QDialog.DialogCode.Accepted:
        # 2. Extract the number from the dialog
        horaire_output(dialog.get_N(),dialog.get_monits(),dialog.get_monits_n_dodo(), dialog.get_n_aides(),dialog.get_n_dortoirs(),dialog.get_n_dodos_ETA())

        #main_window.show()

        QApplication.instance().quit()
    else:
        # If they clicked Cancel or closed the dialog, exit cleanly
        sys.exit(0)
        QApplication.instance().quit()
    
